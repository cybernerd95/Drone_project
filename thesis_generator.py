from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from docx import Document

embeddings = OpenAIEmbeddings()

db = FAISS.load_local(
    "database",
    embeddings,
    allow_dangerous_deserialization=True
)

llm = ChatOpenAI(
    model="gpt-4.1",
    temperature=0.4
)

title = input("Title: ")
methodology = input("Methodology: ")
progress = input("Research Progress: ")

chapters = [
    "Abstract",
    "Introduction",
    "Literature Review",
    "Problem Statement",
    "Research Objectives",
    "Methodology",
    "Implementation",
    "Results",
    "Discussion",
    "Future Work",
    "Conclusion"
]

doc = Document()

doc.add_heading(title, level=1)

for chapter in chapters:

    docs = db.similarity_search(
        chapter + " " + title,
        k=5
    )

    context = "\n\n".join([d.page_content for d in docs])

    prompt = f"""
You are writing an original master's thesis.

Title:
{title}

Methodology:
{methodology}

Research Progress:
{progress}

Reference Material:
{context}

Write a detailed chapter named "{chapter}".

Length:
1000-1500 words.

Write academically.
Do not copy the references.
Do not plagiarize.
"""

    response = llm.invoke(prompt)

    doc.add_heading(chapter, level=2)
    doc.add_paragraph(response.content)

doc.save("Generated_Thesis.docx")

print("Done.")